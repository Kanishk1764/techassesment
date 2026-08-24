import io
import re
import base64
import logging
from typing import Optional, List, Dict, Any, Tuple
from docx import Document
from fastapi import HTTPException
from pypdf import PdfReader

logger = logging.getLogger(__name__)

# Safely attempt PyMuPDF import (prefer modern 'import pymupdf', fallback to 'import fitz')
try:
    import pymupdf as fitz
    HAS_PYMUPDF = True
except Exception:
    try:
        import fitz
        HAS_PYMUPDF = True
    except Exception as e:
        HAS_PYMUPDF = False
        logger.warning(f"PyMuPDF native extension unavailable ({e}). Falling back to pure-python pypdf extractor.")

class ParsedResume:
    def __init__(
        self,
        text: str,
        inferred_name: Optional[str] = None,
        inferred_email: Optional[str] = None,
        inferred_years_experience: Optional[float] = None,
        page_images_base64: Optional[List[str]] = None
    ):
        self.text = text
        self.inferred_name = inferred_name
        self.inferred_email = inferred_email
        self.inferred_years_experience = inferred_years_experience
        self.page_images_base64 = page_images_base64 or []

def parse_resume_bytes(file_bytes: bytes, filename: str, render_images: bool = True) -> ParsedResume:
    ext = filename.split(".")[-1].lower() if "." in filename else ""
    if ext not in ["pdf", "docx", "txt"]:
        raise HTTPException(
            status_code=400,
            detail={"error": {"code": "INVALID_FILE_TYPE", "message": f"Unsupported file type .{ext or 'unknown'}. Only .pdf, .docx, and .txt files are accepted."}}
        )

    if len(file_bytes) > 5 * 1024 * 1024:
        raise HTTPException(
            status_code=400,
            detail={"error": {"code": "FILE_TOO_LARGE", "message": "File exceeds maximum upload size limit (5MB)."}}
        )

    extracted_text = ""
    page_images_base64: List[str] = []

    if ext == "txt":
        try:
            extracted_text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            extracted_text = file_bytes.decode("latin-1", errors="ignore")

    elif ext == "docx":
        try:
            doc = Document(io.BytesIO(file_bytes))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        parts.append(row_text)
            extracted_text = "\n".join(parts)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise HTTPException(
                status_code=400,
                detail={"error": {"code": "DOCX_PARSE_FAILED", "message": "Could not parse .docx file. The document may be corrupted or password-protected."}}
            )

    elif ext == "pdf":
        if HAS_PYMUPDF:
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                text_blocks = []
                for page_idx in range(len(doc)):
                    page = doc[page_idx]
                    page_text = page.get_text("text")
                    if page_text.strip():
                        text_blocks.append(page_text.strip())
                    
                    if render_images and page_idx < 2:
                        try:
                            pix = page.get_pixmap(dpi=150)
                            img_bytes = pix.tobytes("png")
                            b64 = base64.b64encode(img_bytes).decode("utf-8")
                            page_images_base64.append(b64)
                        except Exception:
                            pass

                extracted_text = "\n\n".join(text_blocks)
            except Exception as e:
                logger.error(f"PyMuPDF extraction failed: {e}. Trying pypdf fallback...")
                extracted_text = ""

        # Pure Python pypdf fallback
        if not extracted_text:
            try:
                reader = PdfReader(io.BytesIO(file_bytes))
                pages_text = []
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        pages_text.append(t)
                extracted_text = "\n\n".join(pages_text)
            except Exception as e:
                logger.error(f"pypdf extraction error: {e}")
                raise HTTPException(
                    status_code=400,
                    detail={"error": {"code": "PDF_PARSE_FAILED", "message": "Could not parse PDF file. The file might be corrupted or password-protected."}}
                )

    cleaned_text = extracted_text.replace("\r\n", "\n").strip()
    if not cleaned_text or len(re.sub(r"\s+", "", cleaned_text)) < 20:
        raise HTTPException(
            status_code=400,
            detail={"error": {"code": "NO_EXTRACTABLE_TEXT", "message": "Could not extract text from this document — please upload a text-based PDF, .docx, or .txt file"}}
        )

    inferred_email = _extract_email(cleaned_text)
    inferred_name = _extract_name(cleaned_text, filename)
    inferred_experience = _extract_accurate_experience(cleaned_text)

    return ParsedResume(
        text=cleaned_text,
        inferred_name=inferred_name or _fallback_name(filename),
        inferred_email=inferred_email or f"candidate_{_clean_filename(filename)}@applicant.local",
        inferred_years_experience=inferred_experience,
        page_images_base64=page_images_base64
    )

def _extract_email(text: str) -> Optional[str]:
    match = re.search(r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})", text)
    return match.group(1).lower() if match else None

def _extract_name(text: str, filename: str) -> Optional[str]:
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:4]:
        if (
            3 <= len(line) <= 40
            and "@" not in line
            and "http" not in line
            and not re.search(r"\d", line)
            and not re.search(r"(resume|curriculum|vitae|contact|summary|skills|experience)", line, re.IGNORECASE)
            and re.match(r"^[a-zA-Z\s'-]+$", line)
        ):
            words = line.split()
            if 2 <= len(words) <= 4:
                return line

    return None

def _fallback_name(filename: str) -> str:
    base = re.sub(r"\.(pdf|docx|txt)$", "", filename, flags=re.IGNORECASE).replace("_", " ").replace("-", " ")
    cleaned = re.sub(r"(resume|cv|application|profile)", "", base, flags=re.IGNORECASE).strip()
    if len(cleaned) >= 2 and re.match(r"^[a-zA-Z\s]+$", cleaned):
        return cleaned.title()
    return "Candidate"

def _clean_filename(filename: str) -> str:
    return re.sub(r"[^a-zA-Z0-9]", "", filename.rsplit(".", 1)[0])[:12].lower()

def _extract_accurate_experience(text: str) -> Optional[float]:
    explicit_matches = re.findall(r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)(?:\s+of)?\s+(?:experience|exp)?", text, re.IGNORECASE)
    explicit_exp = 0.0
    for m in explicit_matches:
        try:
            val = float(m)
            if 0 < val <= 40 and val > explicit_exp:
                explicit_exp = val
        except ValueError:
            pass

    year_ranges = re.findall(r"\b(20\d\d|19\d\d)\s*[-–—to]+\s*(20\d\d|present|current|now)\b", text, re.IGNORECASE)
    timeline_years = 0.0
    current_year = 2026

    if year_ranges:
        intervals: List[Tuple[int, int]] = []
        for start_str, end_str in year_ranges:
            try:
                start_yr = int(start_str)
                if end_str.lower() in ["present", "current", "now"]:
                    end_yr = current_year
                else:
                    end_yr = int(end_str)
                if 1980 <= start_yr <= end_yr <= current_year:
                    intervals.append((start_yr, end_yr))
            except ValueError:
                pass

        if intervals:
            intervals.sort(key=lambda x: x[0])
            merged = []
            for start, end in intervals:
                if not merged:
                    merged.append([start, end])
                else:
                    prev = merged[-1]
                    if start <= prev[1]:
                        prev[1] = max(prev[1], end)
                    else:
                        merged.append([start, end])
            timeline_years = sum(end - start for start, end in merged)

    if timeline_years > 0:
        return float(timeline_years)
    if explicit_exp > 0:
        return float(explicit_exp)
    return None
